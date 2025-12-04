"""
Treatment Information Document Service - 待遇情報明示書生成サービス

Generates legally required treatment information documents:
1. 派遣時の待遇情報明示書 (法31条の2第3項) - During dispatch
2. 雇入れ時の待遇情報明示書 (法31条の2第2項) - At hiring
"""
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings


class TreatmentDocumentService:
    """Service for generating treatment information disclosure documents."""

    def __init__(self):
        self.output_dir = Path(settings.PDF_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ========================================
    # UTILITY METHODS
    # ========================================

    def _format_date_japanese(self, d: date) -> str:
        """Format date in Japanese style (2024年12月1日)."""
        if d is None:
            return ""
        return d.strftime("%Y年%m月%d日")

    def _format_date_short(self, d: date) -> str:
        """Short date format (2024.12.1)."""
        if d is None:
            return ""
        return f"{d.year}.{d.month}.{d.day}"

    def _setup_document_a4(self) -> Document:
        """Create a new A4 document with minimal margins."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        return doc

    def _set_cell_font(self, cell, text: str, size: int = 9, bold: bool = False):
        """Set cell text with Japanese font."""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = bold
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_title(self, doc: Document, text: str, size: int = 16):
        """Add centered title."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _add_paragraph(self, doc: Document, text: str, size: int = 10,
                       bold: bool = False, alignment=None) -> None:
        """Add a paragraph with Japanese font."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = bold
        if alignment:
            p.alignment = alignment
        return p

    def _create_bordered_table(self, doc: Document, rows: int, cols: int):
        """Create table with borders."""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        return table

    def _checkbox(self, checked: bool) -> str:
        """Return checkbox character."""
        return "☑" if checked else "☐"

    # ========================================
    # 1. 派遣時の待遇情報明示書 (法31条の2第3項)
    # Treatment Information at Dispatch
    # ========================================

    def generate_haken_ji_taigu_meijisho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 派遣時の待遇情報明示書 (Treatment Info at Dispatch).
        法第31条の2第3項の明示
        派遣元 → 派遣労働者 (when dispatching)

        Required data keys:
        - worker_name: 派遣労働者氏名
        - is_kyotei_taisho: 協定対象派遣労働者か否か
        - kyotei_end_date: 協定の有効期間終了日
        - wage_type: 賃金タイプ (monthly/daily/hourly)
        - base_wage: 基本賃金
        - wage_closing_day: 賃金締切日
        - wage_payment_day: 賃金支払日
        - wage_payment_method: 支払方法
        - has_raise: 昇給の有無
        - has_bonus: 賞与の有無
        - has_retirement_allowance: 退職手当の有無
        - yukyu_after_6months: 6ヶ月継続勤務後の有給日数
        - yukyu_hourly: 時間単位年休の有無
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "派遣時の待遇情報明示書")

        # Recipient line
        p = doc.add_paragraph()
        run = p.add_run(f"\n{data.get('worker_name', '')}　様")
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(12)

        # Issuer info (right aligned)
        p = doc.add_paragraph()
        run = p.add_run(f"（事業所名）{settings.COMPANY_NAME}")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        run = p.add_run(f"（許可番号）{settings.COMPANY_LICENSE_NUMBER}")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Legal basis
        self._add_paragraph(doc, "\n次の条件で労働者派遣を行います(法第31条の2第3項の明示)", 10)

        # ========================================
        # Section 1: 協定対象派遣労働者であるか否か
        # ========================================
        self._add_paragraph(doc, "\n協定対象派遣労働者であるか否か", 11, bold=True)

        is_kyotei = data.get('is_kyotei_taisho', True)
        kyotei_end = self._format_date_japanese(data.get('kyotei_end_date'))

        table = self._create_bordered_table(doc, 2, 1)
        table.rows[0].cells[0].width = Mm(170)

        self._set_cell_font(
            table.rows[0].cells[0],
            f"  {self._checkbox(is_kyotei)} 協定対象派遣労働者である　（当該協定の有効期間の終了　{kyotei_end}　）",
            10
        )
        self._set_cell_font(
            table.rows[1].cells[0],
            f"  {self._checkbox(not is_kyotei)} 協定対象派遣労働者ではない",
            10
        )

        # ========================================
        # Section 2: 賃金に関する事項
        # ========================================
        self._add_paragraph(doc, "\n賃金(退職手当及び臨時に支払われる賃金を除く。)の決定、計算及び支払いの方法、賃金の締切り及び支払いの時期に関する事項", 10, bold=True)

        wage_table = self._create_bordered_table(doc, 6, 4)

        # Row 0: 基本賃金
        self._set_cell_font(wage_table.rows[0].cells[0], "1　基本賃金", 9, True)
        wage_type = data.get('wage_type', 'hourly')
        monthly_check = self._checkbox(wage_type == 'monthly')
        daily_check = self._checkbox(wage_type == 'daily')
        hourly_check = self._checkbox(wage_type == 'hourly')

        self._set_cell_font(wage_table.rows[0].cells[1], f"イ　月給(　　　　円),ロ　日給(　　　　円)", 9)
        self._set_cell_font(wage_table.rows[0].cells[2], "4　賃金締切日(　　)当月", 9)
        self._set_cell_font(wage_table.rows[0].cells[3], "", 9)

        # Row 1
        self._set_cell_font(wage_table.rows[1].cells[0], "", 9)
        base_wage = data.get('base_wage', 0)
        self._set_cell_font(wage_table.rows[1].cells[1], f"ハ　時間給　(　{base_wage:,}　円)", 9)
        self._set_cell_font(wage_table.rows[1].cells[2], f"5　賃金支払日(　　)翌月", 9)
        self._set_cell_font(wage_table.rows[1].cells[3], "", 9)

        # Row 2: 出来高給
        self._set_cell_font(wage_table.rows[2].cells[0], "", 9)
        self._set_cell_font(wage_table.rows[2].cells[1], "ニ　出来高給(基本単価　　　円、保障給　　　円)", 9)
        wage_table.rows[2].cells[2].merge(wage_table.rows[2].cells[3])
        self._set_cell_font(wage_table.rows[2].cells[2], "", 9)

        # Row 3
        self._set_cell_font(wage_table.rows[3].cells[0], "", 9)
        self._set_cell_font(wage_table.rows[3].cells[1], "ホ　その他　(　　　　円)", 9)
        self._set_cell_font(wage_table.rows[3].cells[2], "6　賃金の支払方法(　銀行振込　)", 9)
        self._set_cell_font(wage_table.rows[3].cells[3], "", 9)

        # Row 4
        self._set_cell_font(wage_table.rows[4].cells[0], "", 9)
        self._set_cell_font(wage_table.rows[4].cells[1], "へ　就業規則に規定されている賃金等級等", 9)
        wage_table.rows[4].cells[2].merge(wage_table.rows[4].cells[3])
        self._set_cell_font(wage_table.rows[4].cells[2], "", 9)

        # Row 5: 諸手当
        self._set_cell_font(wage_table.rows[5].cells[0], "2　諸手当の額又は計算方法", 9, True)
        wage_table.rows[5].cells[1].merge(wage_table.rows[5].cells[3])
        allowances = data.get('allowances', [])
        allowance_text = ""
        for i, allow in enumerate(allowances[:4], 1):
            label = ['イ', 'ロ', 'ハ', 'ニ'][i-1]
            allowance_text += f"{label}(　{allow.get('name', '')}手当　　　円　／計算方法:　{allow.get('method', '')}　)\n"
        if not allowance_text:
            allowance_text = "イ(　手当　　　円　／計算方法:　　　)\nロ(　手当　　　円　／計算方法:　　　)"
        self._set_cell_font(wage_table.rows[5].cells[1], allowance_text, 8)

        # ========================================
        # Section 3: 割増賃金率
        # ========================================
        overtime_table = self._create_bordered_table(doc, 3, 2)

        self._set_cell_font(overtime_table.rows[0].cells[0], "3　所定時間外、休日又は深夜労働に対して支払われる割増賃金率", 9, True)
        overtime_table.rows[0].cells[0].merge(overtime_table.rows[0].cells[1])

        ot_rate_normal = data.get('overtime_rate_under_60h', 25)
        ot_rate_over = data.get('overtime_rate_over_60h', 50)
        self._set_cell_font(overtime_table.rows[1].cells[0], f"イ　所定時間外、法定超　月60時間以内(　{ot_rate_normal}　)%", 9)
        self._set_cell_font(overtime_table.rows[1].cells[1], f"月60時間超　(　{ot_rate_over}　)%", 9)

        holiday_rate = data.get('holiday_rate_percent', 35)
        night_rate = data.get('night_rate_percent', 25)
        self._set_cell_font(overtime_table.rows[2].cells[0], f"ロ　休日　法定休日(　{holiday_rate}　)%、法定外休日(　　　)%", 9)
        self._set_cell_font(overtime_table.rows[2].cells[1], f"ハ深夜(　{night_rate}　)%", 9)

        # ========================================
        # Section 4: 昇給・賞与・退職手当の有無
        # ========================================
        self._add_paragraph(doc, "\n昇給・賞与・退職手当の有無", 11, bold=True)

        benefit_table = self._create_bordered_table(doc, 3, 2)

        has_raise = data.get('has_raise', True)
        has_bonus = data.get('has_bonus', False)
        has_retirement = data.get('has_retirement_allowance', False)

        self._set_cell_font(benefit_table.rows[0].cells[0], f"・昇　給　（{self._checkbox(has_raise)}有　（時期、金額等", 9)
        self._set_cell_font(benefit_table.rows[0].cells[1], f"）,{self._checkbox(not has_raise)}無　）", 9)

        self._set_cell_font(benefit_table.rows[1].cells[0], f"・賞　与　（{self._checkbox(has_bonus)}有　（時期、金額等", 9)
        self._set_cell_font(benefit_table.rows[1].cells[1], f"）,{self._checkbox(not has_bonus)}無　）", 9)

        self._set_cell_font(benefit_table.rows[2].cells[0], f"・退職手当（{self._checkbox(has_retirement)}有　（時期、金額等", 9)
        self._set_cell_font(benefit_table.rows[2].cells[1], f"）,{self._checkbox(not has_retirement)}無　）", 9)

        # ========================================
        # Section 5: 休暇に関する事項
        # ========================================
        self._add_paragraph(doc, "\n休暇に関する事項", 11, bold=True)

        leave_table = self._create_bordered_table(doc, 3, 1)

        yukyu_days = data.get('yukyu_after_6months', 10)
        has_hourly = data.get('yukyu_hourly', False)

        self._set_cell_font(leave_table.rows[0].cells[0],
            f"1　年次有給休暇　6か月継続勤務した場合→{yukyu_days}日　　継続勤務6か月以内の年次有給休暇　（{self._checkbox(has_hourly)}有・{self._checkbox(not has_hourly)}無）→　ヶ月経過で　日", 9)
        self._set_cell_font(leave_table.rows[0].cells[0],
            f"1　年次有給休暇　6か月継続勤務した場合→{yukyu_days}日\n　時間単位年休（{self._checkbox(has_hourly)}有・{self._checkbox(not has_hourly)}無）", 9)

        has_substitute = data.get('has_substitute_leave', False)
        self._set_cell_font(leave_table.rows[1].cells[0],
            f"2　代替休暇（{self._checkbox(has_substitute)}有・{self._checkbox(not has_substitute)}無）", 9)

        other_leave = data.get('other_leave', '')
        self._set_cell_font(leave_table.rows[2].cells[0],
            f"3　その他の休暇　有給(　{other_leave}　)\n　　　　　　　　無給(　　　　　　　　　　　　　　　　）", 9)

        # Footer
        self._add_paragraph(doc, "\n○詳細は、就業規則第3章　第2節　休暇　第43条〜第45条", 9)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 2. 雇入れ時の待遇情報明示書 (法31条の2第2項)
    # Treatment Information at Hiring
    # ========================================

    def generate_yatoire_ji_taigu_meijisho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 雇入れ時の待遇情報明示書 (Treatment Info at Hiring).
        法第31条の2第2項の明示
        派遣元 → 派遣労働者 (when hiring)

        Required data keys:
        - document_date: 書類作成日
        - worker_name: 派遣労働者氏名
        - is_kyotei_taisho: 協定対象派遣労働者か否か
        - kyotei_end_date: 協定の有効期間終了日
        - has_raise: 昇給の有無
        - has_bonus: 賞与の有無
        - has_retirement_allowance: 退職手当の有無
        - hakensaki_complaint_dept: 派遣先苦情処理担当部署
        - hakensaki_complaint_name: 派遣先苦情処理担当者名
        - hakensaki_complaint_phone: 派遣先苦情処理担当電話
        - hakenmoto_complaint_dept: 派遣元苦情処理担当部署
        - hakenmoto_complaint_name: 派遣元苦情処理担当者名
        - hakenmoto_complaint_phone: 派遣元苦情処理担当電話
        """
        doc = self._setup_document_a4()

        # Document date (right aligned)
        doc_date = data.get('document_date', date.today())
        p = doc.add_paragraph()
        run = p.add_run(self._format_date_japanese(doc_date))
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Title
        self._add_title(doc, "雇入れ時の待遇情報明示書")

        # Subtitle
        p = doc.add_paragraph()
        run = p.add_run("(法31条の2第2項)")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Recipient line
        p = doc.add_paragraph()
        run = p.add_run(f"\n\n{data.get('worker_name', '')}　殿")
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(12)

        # Issuer info (right aligned)
        p = doc.add_paragraph()
        run = p.add_run(f"\n（事業所名）{settings.COMPANY_NAME}")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        run = p.add_run(f"（許可番号）{settings.COMPANY_LICENSE_NUMBER}")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ========================================
        # Main content table
        # ========================================
        main_table = self._create_bordered_table(doc, 4, 2)

        for row in main_table.rows:
            row.cells[0].width = Mm(40)
            row.cells[1].width = Mm(130)

        # Row 0: 協定対象派遣労働者であるか否か
        self._set_cell_font(main_table.rows[0].cells[0], "協定対象派遣\n労働者であるか\n否か", 9, True)

        is_kyotei = data.get('is_kyotei_taisho', True)
        kyotei_end = self._format_date_japanese(data.get('kyotei_end_date'))

        kyotei_text = f"{self._checkbox(is_kyotei)}協定対象派遣労働者である（当該協定の有効期間の終了日: {kyotei_end}）\n\n{self._checkbox(not is_kyotei)}協定対象派遣労働者ではない"
        self._set_cell_font(main_table.rows[0].cells[1], kyotei_text, 10)

        # Row 1: 昇給・賞与・退職手当の有無
        self._set_cell_font(main_table.rows[1].cells[0], "昇給・賞与・退職\n手当の有無", 9, True)

        has_raise = data.get('has_raise', True)
        has_bonus = data.get('has_bonus', False)
        has_retirement = data.get('has_retirement_allowance', False)

        benefit_text = f"・昇　給　（　{self._checkbox(has_raise)}有　（時期、金額等　　　　　　　　　　　　）　,　{self._checkbox(not has_raise)}無　）\n"
        benefit_text += f"・賞　与　（　{self._checkbox(has_bonus)}有　（時期、金額等　　　　　　　　　　　　）　,　{self._checkbox(not has_bonus)}無　）\n"
        benefit_text += f"・退職手当（　{self._checkbox(has_retirement)}有　（時期、金額等　　　　　　　　　　　　）　,　{self._checkbox(not has_retirement)}無　）"

        self._set_cell_font(main_table.rows[1].cells[1], benefit_text, 10)

        # Row 2: 苦情の申し出先・処理方法・連携体制 (1)
        # Get complaint handler info
        saki_dept = data.get('hakensaki_complaint_dept', '')
        saki_name = data.get('hakensaki_complaint_name', '')
        saki_phone = data.get('hakensaki_complaint_phone', '')
        moto_dept = data.get('hakenmoto_complaint_dept', settings.DISPATCH_COMPLAINT_DEPARTMENT)
        moto_name = data.get('hakenmoto_complaint_name', settings.DISPATCH_COMPLAINT_NAME)
        moto_phone = data.get('hakenmoto_complaint_phone', settings.DISPATCH_COMPLAINT_PHONE)

        self._set_cell_font(main_table.rows[2].cells[0], "(1)苦情の申し出先\n・処理方法\n・連携体制", 9, True)

        complaint_text = f"派遣先(部署)　{saki_dept}　　　　　　{saki_name}　　TEL　{saki_phone}\n"
        complaint_text += f"派遣元(部署)　{moto_dept}　　取締役 部長　{moto_name}　　TEL　{moto_phone}"

        self._set_cell_font(main_table.rows[2].cells[1], complaint_text, 10)

        # Row 3: 苦情処理方法 (2)
        self._set_cell_font(main_table.rows[3].cells[0], "(2)苦情処理方法", 9, True)

        # Legal text for complaint handling procedure
        procedure_text = """①派遣先における(1)記載の者が苦情の申し出を受けた時は、直ちに派遣先責任者へ連絡し当該派遣先責任者が中心となり誠意を持って遅滞なく当該苦情処理の適切かつ迅速な処理を図り、その結果について必ず派遣労働者に通知します。

②派遣元における(1)の記載の者が苦情の申し出を受けた時は、直ちに派遣元責任者へ連絡し当該派遣元責任者が中心となり誠意をもって遅滞なく当該苦情の適切かつ迅速な処理を図り、その結果について必ず派遣労働者に通知します。

③派遣先及び派遣元は、自らでその解決が容易であり即日に処理した苦情の他は、相互に遅滞なく通知するとともに、密接に連絡調整を行いつつ、その解決をはかることとする。"""

        self._set_cell_font(main_table.rows[3].cells[1], procedure_text, 8)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
