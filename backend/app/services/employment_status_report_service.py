"""
Employment Status Report Service - 就業状況報告書生成サービス

Generates 就業状況報告書 (Employment Status Report) documents.
This document reports the working status of dispatched employees to the client company.
"""
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings


class EmploymentStatusReportService:
    """Service for generating Employment Status Report (就業状況報告書) documents."""

    def __init__(self):
        self.output_dir = Path(settings.PDF_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ========================================
    # UTILITY METHODS
    # ========================================

    def _format_date_japanese(self, d: date) -> str:
        """Format date in Japanese style (2024年12月1日)."""
        if d is None:
            return ""
        return d.strftime("%Y年%m月%d日")

    def _format_date_slash(self, d: date) -> str:
        """Format date with slashes (2024/12/01)."""
        if d is None:
            return ""
        return d.strftime("%Y/%m/%d")

    def _setup_document_a4(self) -> Document:
        """Create a new A4 document with minimal margins."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        return doc

    def _set_cell_font(self, cell, text: str, size: int = 9, bold: bool = False):
        """Set cell text with Japanese font."""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = bold
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_title(self, doc: Document, text: str, size: int = 16):
        """Add centered title."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _add_paragraph(self, doc: Document, text: str, size: int = 10,
                       bold: bool = False, alignment=None) -> None:
        """Add a paragraph with Japanese font."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = bold
        if alignment:
            p.alignment = alignment
        return p

    def _create_bordered_table(self, doc: Document, rows: int, cols: int):
        """Create table with borders."""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        return table

    # ========================================
    # 就業状況報告書 (Employment Status Report)
    # ========================================

    def generate_shugyo_jokyo_hokokusho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 就業状況報告書 (Employment Status Report).
        派遣先 → 派遣元 (Monthly report of dispatched workers status)

        Required data keys:
        - factory_name: 派遣先名称 (client company)
        - factory_address: 派遣先所在地
        - start_date: 報告期間開始日
        - end_date: 報告期間終了日
        - break_time_minutes: 休憩時間 (minutes, default 80)
        - employees: List of employee data, each containing:
            - employee_number: 番号
            - name: 氏名
            - work_content: 従事した業務内容
            - organization_unit: 組織単位
            - department: 部署名
            - responsibility_level: 責任の程度 (default: 付与される権限なし)
        """
        doc = self._setup_document_a4()

        # ========================================
        # Header Section
        # ========================================

        # Recipient line (派遣元会社)
        p = doc.add_paragraph()
        run = p.add_run(f"{settings.COMPANY_NAME}　御中")
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(12)

        # Title
        self._add_title(doc, "就業状況報告書")

        # Period line
        start_date = data.get('start_date', date.today())
        end_date = data.get('end_date', date.today())

        period_text = f"{self._format_date_slash(start_date)}　〜　{self._format_date_japanese(end_date)}　の就業状況を以下の通りご報告します"
        self._add_paragraph(doc, period_text, 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Break time note
        break_minutes = data.get('break_time_minutes', 80)
        break_note = f"尚、休憩時間については、全作業者{break_minutes}分となっております。"
        self._add_paragraph(doc, break_note, 9)

        # ========================================
        # Dispatch Location Section
        # ========================================
        self._add_paragraph(doc, "\n派遣就業した場所", 11, bold=True)

        location_table = self._create_bordered_table(doc, 2, 2)

        # Set column widths
        for row in location_table.rows:
            row.cells[0].width = Mm(25)
            row.cells[1].width = Mm(150)

        factory_name = data.get('factory_name', '')
        factory_address = data.get('factory_address', '')

        self._set_cell_font(location_table.rows[0].cells[0], "(名称)", 9, True)
        self._set_cell_font(location_table.rows[0].cells[1], factory_name, 10)
        self._set_cell_font(location_table.rows[1].cells[0], "(所在地)", 9, True)
        self._set_cell_font(location_table.rows[1].cells[1], factory_address, 10)

        # ========================================
        # Employee List Table
        # ========================================
        self._add_paragraph(doc, "", 6)  # Spacer

        employees = data.get('employees', [])

        # Create table with header + employee rows
        num_rows = len(employees) + 1  # +1 for header
        employee_table = self._create_bordered_table(doc, num_rows, 6)

        # Define column widths (in Mm)
        col_widths = [10, 35, 55, 20, 20, 40]  # Total ~180mm

        # Set header row
        headers = [
            "No",
            "氏名",
            "従事した業務内容",
            "組織単位",
            "部署名",
            "派遣労働者が従事する\n業務に伴う責任の程度"
        ]

        for idx, header in enumerate(headers):
            cell = employee_table.rows[0].cells[idx]
            cell.width = Mm(col_widths[idx])
            self._set_cell_font(cell, header, 8, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fill employee data
        for row_idx, emp in enumerate(employees, 1):
            if row_idx >= num_rows:
                break

            row = employee_table.rows[row_idx]

            # No (sequential number)
            self._set_cell_font(row.cells[0], str(emp.get('employee_number', row_idx)), 9)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 氏名
            self._set_cell_font(row.cells[1], emp.get('name', ''), 9)

            # 従事した業務内容
            self._set_cell_font(row.cells[2], emp.get('work_content', ''), 8)

            # 組織単位
            org_unit = emp.get('organization_unit', '0')
            self._set_cell_font(row.cells[3], str(org_unit) if org_unit else '0', 9)
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 部署名
            dept = emp.get('department', '0')
            self._set_cell_font(row.cells[4], str(dept) if dept else '0', 9)
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 責任の程度
            responsibility = emp.get('responsibility_level', '付与される権限なし')
            self._set_cell_font(row.cells[5], responsibility, 8)

        # ========================================
        # Footer Section
        # ========================================

        # Add space for signature if needed
        self._add_paragraph(doc, "\n", 10)

        # Client company signature area
        p = doc.add_paragraph()
        run = p.add_run("上記の通り報告します。")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)

        # Date and signature line
        p = doc.add_paragraph()
        run = p.add_run(f"\n　　年　　月　　日")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        sender_name = data.get('sender_company_name', '')
        run = p.add_run(f"\n　　　　　　　　　　　　　　　　　　　　　　　　（派遣先）{sender_name}")
        run.font.name = "MS Gothic"
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_shugyo_jokyo_hokokusho_from_db(
        self,
        db_session,
        factory_id: int,
        start_date: date,
        end_date: date
    ) -> bytes:
        """
        Generate 就業状況報告書 by querying data from database.

        This method queries the database for:
        - Factory information
        - Active employees assigned to the factory
        - Employee information

        Args:
            db_session: SQLAlchemy database session
            factory_id: ID of the factory (派遣先)
            start_date: Report period start date
            end_date: Report period end date

        Returns:
            bytes: Generated DOCX document
        """
        from app.models.factory import Factory
        from app.models.employee import Employee
        from sqlalchemy import select, and_, or_

        # Get factory info
        factory_result = db_session.execute(
            select(Factory).where(Factory.id == factory_id)
        )
        factory = factory_result.scalar_one_or_none()

        if not factory:
            raise ValueError(f"Factory with id {factory_id} not found")

        # Get employees assigned to this factory
        # Employees have a factory_id field linking them to their current factory
        employees_result = db_session.execute(
            select(Employee).where(
                and_(
                    Employee.factory_id == factory_id,
                    Employee.status == 'active'  # Only active employees
                )
            ).order_by(Employee.employee_number)
        )

        employees_data = []
        for idx, employee in enumerate(employees_result.scalars().all(), 1):
            # Get work content from factory settings or default
            work_content = '機械オペレーター及び機械メンテナンス他付随する業務'

            employees_data.append({
                'employee_number': idx,
                'name': employee.full_name_kanji or employee.full_name_romaji or '',
                'work_content': work_content,
                'organization_unit': '0',
                'department': employee.department or '0',
                'responsibility_level': '付与される権限なし'
            })

        # Prepare data for document generation
        data = {
            'factory_name': f"{factory.company_name} {factory.plant_name or ''}".strip(),
            'factory_address': factory.plant_address or factory.company_address or '',
            'start_date': start_date,
            'end_date': end_date,
            'break_time_minutes': factory.break_minutes or 80,
            'employees': employees_data,
            'sender_company_name': factory.company_name or ''
        }

        return self.generate_shugyo_jokyo_hokokusho(data)
